"""File parsing utilities for PDF / DOCX / Markdown / video subtitles."""
from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable

from utils.logger import logger


def parse_pdf(path: str | Path) -> str:
    """Extract text from a PDF using pdfplumber for table-friendly extraction."""
    import pdfplumber

    path = Path(path)
    text_parts: list[str] = []
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            try:
                text_parts.append(page.extract_text() or "")
            except Exception as exc:  # pragma: no cover
                logger.warning(f"PDF page {i} extraction failed: {exc}")
    return "\n\n".join(p for p in text_parts if p.strip())


def parse_docx(path: str | Path) -> str:
    """Extract text + tables from .docx files."""
    from docx import Document

    doc = Document(str(path))
    parts: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            parts.append(" | ".join(cells))
    return "\n".join(parts)


def parse_markdown(path: str | Path) -> str:
    """Read a markdown file and return its raw text."""
    return Path(path).read_text(encoding="utf-8", errors="ignore")


def parse_srt(path: str | Path) -> str:
    """Strip timestamps from .srt subtitles."""
    raw = Path(path).read_text(encoding="utf-8", errors="ignore")
    cleaned = re.sub(
        r"\d+\n\d{2}:\d{2}:\d{2}[.,]\d{3} --> \d{2}:\d{2}:\d{2}[.,]\d{3}\n",
        "",
        raw,
    )
    return cleaned.strip()


def parse_vtt(path: str | Path) -> str:
    """Strip metadata from .vtt subtitles."""
    lines: Iterable[str] = Path(path).read_text(encoding="utf-8", errors="ignore").splitlines()
    out: list[str] = []
    for line in lines:
        if not line.strip() or line.startswith(("WEBVTT", "NOTE")) or "-->" in line:
            continue
        out.append(line)
    return "\n".join(out)


PARSERS = {
    ".pdf": parse_pdf,
    ".docx": parse_docx,
    ".md": parse_markdown,
    ".markdown": parse_markdown,
    ".txt": parse_markdown,
    ".srt": parse_srt,
    ".vtt": parse_vtt,
}


def parse_file(path: str | Path) -> str:
    """Dispatch to the right parser based on file extension."""
    p = Path(path)
    ext = p.suffix.lower()
    if ext not in PARSERS:
        raise ValueError(f"Unsupported file type: {ext}")
    logger.info(f"Parsing {p.name} ({ext})")
    return PARSERS[ext](p)
